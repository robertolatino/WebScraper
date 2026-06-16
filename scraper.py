import os
import time
import html
from playwright.sync_api import sync_playwright
from bs4 import BeautifulSoup
from docx import Document

# Mantengo tus funciones de ayuda intactas
def leer_codigos_desde_txt(nombre_archivo="codigos.txt"):
    """Lee el archivo .txt y devuelve una lista con los códigos limpios."""
    if not os.path.exists(nombre_archivo):
        return []
    with open(nombre_archivo, "r", encoding="utf-8") as f:
        codigos = [linea.strip() for linea in f.readlines() if linea.strip()]
    return codigos

def agregar_html_a_docx(html_text, doc, codigo):
    """Convierte HTML a un párrafo de Word manteniendo negritas y cursivas."""
    doc.add_paragraph(codigo, style='Heading 2')
    p = doc.add_paragraph("Enunciado: ")

    if not html_text or html_text == "ERROR":
        p.add_run(html_text)
        doc.add_paragraph() 
        return

    texto_decodificado = html.unescape(html_text)
    soup = BeautifulSoup(texto_decodificado, "html.parser")

    for elemento in soup.descendants:
        if elemento.name is None: 
            texto = str(elemento)
            if texto == '\n': 
                continue
            run = p.add_run(texto)
            padre = elemento.parent
            while padre and padre.name != '[document]':
                if padre.name in ['b', 'strong']: run.bold = True
                if padre.name in ['i', 'em']: run.italic = True
                if padre.name in ['u']: run.underline = True
                padre = padre.parent
    doc.add_paragraph() 

# --- ESTA ES LA FUNCIÓN QUE CONECTA CON APP.PY ---
def ejecutar_extraccion_enunciados(url_base, usuario, contrasena, lista_codigos, progreso_st=None):
    LISTA_CONTENIDOS = lista_codigos
    
    if not LISTA_CONTENIDOS:
        if progreso_st:
            progreso_st.update(label="No hay códigos para procesar.", state="error")
        return []

    with sync_playwright() as p:
        # headless=True para que corra oculto en la aplicación web
        browser = p.firefox.launch(headless=True)  
        context = browser.new_context()
        page = context.new_page()
        
        def aplicar_filtro_global():
            try:
                page.locator('.input-search-filter__button').first.click()
                page.wait_for_timeout(1000)
                page.locator('[id="Tipo de contenido"]').first.click()
                page.wait_for_timeout(1000)
                page.locator('label:has-text("Question")').first.click()
                page.wait_for_timeout(500)
                page.locator('button[aria-label="Aplicar filtros"]').first.click()
                page.wait_for_load_state("networkidle")
                page.wait_for_timeout(2000)
            except Exception:
                pass

        # Conectando
        if progreso_st: progreso_st.update(label="Conectando a la plataforma e iniciando sesión...", state="running")
        page.goto(url_base + "/auth/login")
        
        page.get_by_placeholder("Nombre de usuario").fill(usuario)
        page.get_by_placeholder("Contraseña").fill(contrasena)
        page.get_by_role("button", name="Iniciar sesión").click()
        page.wait_for_load_state("networkidle")
        
        # Navegando
        if progreso_st: progreso_st.update(label="Navegando a la sección de Contenidos...", state="running")
        page.locator('div[aria-label="Contenidos"]').first.click()
        page.wait_for_timeout(1000)
        
        page.locator('.wrapper-list-menu__item a[href="/contents"]').first.click()
        page.wait_for_load_state("networkidle")
        page.wait_for_timeout(1500)
        
        aplicar_filtro_global()

        resultados = []
        indice_actual = 0
        intentos_filtro = 0
        
        while indice_actual < len(LISTA_CONTENIDOS):
            codigo = LISTA_CONTENIDOS[indice_actual]
            numero_item = indice_actual + 1
            
            # --- COMUNICACIÓN CON LA WEB ---
            # Esto es clave para que la página web muestre por qué código va en tiempo real
            if progreso_st: 
                progreso_st.update(label=f"Procesando [{numero_item}/{len(LISTA_CONTENIDOS)}]: {codigo}...", state="running")
            
            try:
                buscador = page.locator('input[data-testid="search"]')
                buscador.fill("") 
                buscador.fill(codigo)
                buscador.press("Enter")
                
                enlace_resultado = page.locator(f'a:text-is("{codigo}")').first
                enlace_resultado.wait_for(state="visible", timeout=10000)
                
                fila = page.locator(f'tr:has(a:text-is("{codigo}"))').first
                es_question = fila.locator('td:text-is("Question")').is_visible()
                
                if not es_question:
                    if intentos_filtro < 1:
                        aplicar_filtro_global()
                        intentos_filtro += 1
                        continue 
                    else:
                        resultados.append((codigo, "[NO ES QUESTION]"))
                        intentos_filtro = 0
                        indice_actual += 1
                        continue
                
                intentos_filtro = 0 
                
                enlace_resultado.click()
                page.wait_for_load_state("networkidle")
                page.wait_for_timeout(1000)
                
                btn_editor = page.get_by_role("button", name="Editor", exact=False).first
                
                try:
                    btn_editor.wait_for(state="visible", timeout=8000)
                    btn_editor.click()
                    page.wait_for_load_state("networkidle")
                    page.wait_for_timeout(2000) 
                    
                    bloque_enunciado = page.locator('div[data-id="stimulus"]')
                    
                    try:
                        bloque_enunciado.wait_for(state="visible", timeout=8000)
                        textarea_oculto = bloque_enunciado.locator('textarea').first
                        textarea_oculto.wait_for(state="attached", timeout=3000)
                        
                        descripcion_html = textarea_oculto.evaluate("el => el.value")
                        resultados.append((codigo, descripcion_html))
                        
                    except Exception:
                        resultados.append((codigo, "[SIN ENUNCIADO EN EL EDITOR]"))

                except Exception:
                    resultados.append((codigo, "[SIN PESTAÑA EDITOR]"))

                btn_volver = page.locator('button[aria-label="Volver"]').first
                btn_volver.wait_for(state="visible", timeout=4000)
                btn_volver.click()
                page.wait_for_load_state("networkidle")
                page.wait_for_timeout(1500)
                
                indice_actual += 1
                
            except Exception as e:
                intentos_filtro = 0
                indice_actual += 1
        
        browser.close()
        return resultados


# Dejo tu código standalone intacto por si alguna vez quieres correr scraper.py directamente desde la terminal
def iniciar_automatizacion():
    # Variables locales para modo standalone
    URL_LOCAL = "https://publisher.edelvivesdigitalplus.com"
    USUARIO_LOCAL = "roberto.latino@edelvives.es"
    CONTRASEÑA_LOCAL = "Masaya2022$"
    
    LISTA_CONTENIDOS = leer_codigos_desde_txt("codigos.txt")
    if not LISTA_CONTENIDOS:
        print("Finalizando ejecución porque no hay códigos para procesar.")
        return

    resultados = ejecutar_extraccion_enunciados(URL_LOCAL, USUARIO_LOCAL, CONTRASEÑA_LOCAL, LISTA_CONTENIDOS)
    
    if not resultados:
        return
    
    print("\n[3] Generando documento de Word...")
    doc = Document()
    doc.add_heading('Enunciados de Actividades', 0)
    for cod, desc_html in resultados:
        agregar_html_a_docx(desc_html, doc, cod)
        
    doc.save("enunciados_formateados.docx")
    print("[FIN] Prueba completada. Revisa el archivo 'enunciados_formateados.docx' en tu carpeta.")

if __name__ == "__main__":
    iniciar_automatizacion()